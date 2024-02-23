from docx import Document

def find_tables_with_titles(docx_file):
    document = Document(docx_file)
    tables_with_titles = []

    paragraphs = document.paragraphs
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]

        if paragraph.style.name == 'Caption':
            print(f"Found a paragraph with 'Caption' style: {paragraph.text}")  # Debug print
            if i + 1 < len(document.tables):
                next_table = document.tables[i]
                tables_with_titles.append((paragraph.text, next_table))

    return tables_with_titles

if __name__ == "__main__":
    docx_file_path = "./multi+parent/rep.docx"  # Replace with your .docx file path
    tables_with_titles = find_tables_with_titles(docx_file_path)
    print(f"Found {len(tables_with_titles)} tables with titles")  # Debug print
    for title, table in tables_with_titles:
        print("Title:", title)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
            print()
        print("=" * 20)