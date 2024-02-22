from docx import Document

def extract_text_from_docx(docx_file):
    """
    Extracts text from a .docx file including headings and paragraphs.
    
    Args:
    - docx_file: Path to the .docx file
    
    Returns:
    - text: A string containing extracted headings and paragraphs
    """
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.hyperlinks:
            continue
        if not (paragraph.text.startswith('Table') or paragraph.text.startswith('Figure')):
            text += f"{paragraph.text}\n"
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            text += f"{paragraph.text}\n"
    return text

def write_to_text_file(text, output_file):
    """
    Writes extracted text to a text file.
    
    Args:
    - text: The text to be written
    - output_file: Path to the output text file
    """
    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(text)

def main():
    # Path to the input .docx file
    docx_file_path = "./recursive/rep.docx"
    
    # Path to the output text file
    output_file_path = "./remove intro docs/output.txt"
    
    # Extract text from the .docx file
    extracted_text = extract_text_from_docx(docx_file_path)
    
    # Write extracted text to a text file
    write_to_text_file(extracted_text, output_file_path)
    print("Extraction and writing completed successfully.")

if __name__ == "__main__":
    main()
